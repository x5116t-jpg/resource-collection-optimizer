"""
Extract text content from .docx files for LaTeX report generation
"""
import sys
from docx import Document

def extract_docx_content(filepath):
    """Extract all text from a .docx file"""
    try:
        doc = Document(filepath)
        content = []

        # Extract all paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                content.append(para.text)

        # Extract tables
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells]
                content.append(" | ".join(row_text))

        return "\n".join(content)
    except Exception as e:
        return f"Error reading file: {str(e)}"

if __name__ == "__main__":
    import io
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

    if len(sys.argv) < 2:
        print("Usage: python extract_docx.py <filepath>")
        sys.exit(1)

    filepath = sys.argv[1]
    content = extract_docx_content(filepath)
    print(content)
